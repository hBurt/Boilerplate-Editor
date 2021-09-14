import glob
import json
import os
import re
import subprocess
import tkinter as tk

from datetime import date
from pathlib import Path
from docx import Document

# constants

# gui screen start position assumes monitors are setup in parallel(left to right) and the primary monitor is on the left
# gui screen start position(1 = middle of first screen, 4 = middle of second screen)
gui_start_position = 4

# this truncates n number of lines from the output when copying to the clipboard(does not effect pdf output)
truncate_lines_before = 5
truncate_lines_after = 6

converted_file_base_name = "HarrisonStewardCoverLetter_"
config_file_name = "config.json"

docx_template_path = "C:\\Users\\Harrison\\Documents\\HarrisonStewardCoverLetter_Template.docx"
libre_office_swriter_path = "C:\\Program Files\\LibreOffice\\program\\swriter.exe"
pdf_output_directory = "C:\\Users\\Harrison\\Documents\\"

default_curr_date = re.compile("zCurrentDatez")
default_company_or_hiring_manager_name = re.compile("zCompanyOrHiringManagerNamez")
default_city_state = re.compile("zCityStatez")
default_company = re.compile("zCompanyz")
default_job_title = re.compile("zJobTitlez")


def load_data(json_file_name):
    with open(json_file_name, encoding='utf-8') as file:
        data = json.load(file)
    return data


def add_custom_link(custom_title, custom_text, toolbar_frame, popup_diag_frame):
    new_custom_button = tk.Button(toolbar_frame, text=custom_title)
    new_custom_button.pack(side=tk.LEFT, padx=2, pady=2)

    if not Path(config_file_name).exists():
        data = {
            "customButtons": [
                {
                    "title": custom_title,
                    "text": custom_text
                }
            ]
        }

        with open(config_file_name, 'w', encoding='utf-8') as file:
            json.dump(data, file)
    else:
        loaded_data = load_data(config_file_name)
        new_button = {"title": custom_title, "text": custom_text}
        loaded_data['customButtons'].append(new_button)

        with open(config_file_name, 'w', encoding='utf-8') as file:
            json.dump(loaded_data, file)

    popup_diag_frame.destroy()


def custom_link_popup(toolbar_frame):
    popup_diag_width = 380
    popup_diag_height = 112

    popup_diag = tk.Toplevel()
    popup_diag.iconphoto(False, tk.PhotoImage(file='scroll_x64.png'))
    popup_diag.title("Custom Link")
    popup_diag.geometry(str(popup_diag_width) + "x" + str(popup_diag_height))

    label_custom_link_title = tk.Label(popup_diag, text="Custom link name: ")
    label_custom_link_text = tk.Label(popup_diag, text="Custom link text: ")

    entry_custom_link_title = tk.Entry(popup_diag, width=36)
    entry_custom_link_text = tk.Entry(popup_diag, width=36)

    button_add_custom_link = tk.Button(popup_diag,
                                       text="Add Custom Link",
                                       command=lambda: add_custom_link(entry_custom_link_title.get(),
                                                                       entry_custom_link_text.get(),
                                                                       toolbar_frame,
                                                                       popup_diag))

    label_custom_link_title.grid(row=0, column=0, pady=2, padx=20)
    label_custom_link_text.grid(row=1, column=0, pady=2, padx=20)

    entry_custom_link_title.grid(row=0, column=1, pady=10)
    entry_custom_link_text.grid(row=1, column=1, pady=10)

    button_add_custom_link.grid(row=2, column=0, padx=(20, 0), columnspan=2, sticky=tk.N + tk.S + tk.W + tk.E)

    # calculate coordination of screen and window form
    popup_x_center_main_window = int(window.winfo_x() + ((window.winfo_width() - popup_diag_width) / 2))
    popup_y_center_main_window = int(window.winfo_y() + ((window.winfo_height() - popup_diag_height) / 2))

    # Write following format for center screen
    popup_diag.geometry("+{}+{}".format(popup_x_center_main_window, popup_y_center_main_window))

    entry_custom_link_title.focus_set()

    popup_diag.mainloop()


# https://github.com/python-openxml/python-docx/issues/30#issuecomment-879593691
def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


def convert_call_back(job_title,
                      company_name,
                      city_state,
                      hiring_manager,
                      is_using_hiring_manager_name,
                      is_remote,
                      is_copy_txt_to_clipboard):
    file_name = converted_file_base_name + company_name + ".docx"
    document = Document(docx=docx_template_path)

    today = date.today()

    # replace default words(city/state, jobTitle, companyName)
    for p in document.paragraphs:

        paragraph_replace_text(p, default_curr_date, today.strftime("%B %d, %Y"))

        if is_using_hiring_manager_name == 1:
            paragraph_replace_text(p, default_company_or_hiring_manager_name, hiring_manager)
        else:
            paragraph_replace_text(p, default_company_or_hiring_manager_name, company_name)

        if is_remote == 1:
            paragraph_replace_text(p, default_city_state, city_state + "\nRemote")
        else:
            paragraph_replace_text(p, default_city_state, city_state)

        paragraph_replace_text(p, default_company, company_name)
        paragraph_replace_text(p, default_job_title, job_title)

    document.save(file_name)

    if is_copy_txt_to_clipboard == 1:
        # create the txt file from our *.docx file
        with open("output.txt", "w") as text_file:
            for p in document.paragraphs:
                tmp_stripped_string = re.sub(r'\t', '', p.text)
                text_file.write(tmp_stripped_string)
                text_file.write('\n')

        # read the text file to prepare for truncation
        with open('output.txt', 'r') as fin:
            data = fin.read().splitlines(True)

        # write the data back to the text file
        with open('output.txt', 'w') as fout:
            fout.writelines(data[truncate_lines_before:-truncate_lines_after])

        # copy text file to clipboard
        os.system("clip < output.txt ")

        # delete text file
        os.remove("output.txt")

    else:
        subprocess.run(
            [libre_office_swriter_path, "--headless", "--convert-to", "pdf",
             file_name, "--outdir", pdf_output_directory])

    os.remove(file_name)


def set_hiring_manager_entry_state(state):
    if state == 1:
        entry_hiring_manager_name.config(state='normal')
    else:
        entry_hiring_manager_name.config(state='disabled')


# Start GUI
tk_width = 400
tk_height = 290

window = tk.Tk()
window.iconphoto(False, tk.PhotoImage(file='scroll_x64.png'))
window.title("Boilerplate Editor")
window.geometry(str(tk_width) + "x" + str(tk_height))

is_remote_val = tk.IntVar()
is_remote_val.set(1)

is_copy_to_clipboard_val = tk.IntVar()
is_copy_to_clipboard_val.set(0)

is_hiring_manager_entry_enabled = tk.IntVar()
is_hiring_manager_entry_enabled.set(0)

label_job_title = tk.Label(window, text="Job Title: ")
label_company_name = tk.Label(window, text="Company Name: ")
label_city_state = tk.Label(window, text="City/State: ")
label_hiring_manager_name = tk.Label(window, text="Hiring Manager: ")

entry_job_title = tk.Entry(window, width=40)
entry_company_name = tk.Entry(window, width=40)
entry_city_state = tk.Entry(window, width=40)
entry_hiring_manager_name = tk.Entry(window, width=40, state='disabled')

check_is_remote = tk.Checkbutton(window, text="Is Remote? ", variable=is_remote_val)
check_is_txt_file = tk.Checkbutton(window, text="Copy as text to clipboard", variable=is_copy_to_clipboard_val)
check_is_hiring_manager = tk.Checkbutton(window,
                                         text="Use hiring manager's name",
                                         variable=is_hiring_manager_entry_enabled,
                                         command=lambda:
                                         set_hiring_manager_entry_state(is_hiring_manager_entry_enabled.get()))

button_convert = tk.Button(window, text="Convert",
                           command=lambda: convert_call_back(entry_job_title.get(),
                                                             entry_company_name.get(),
                                                             entry_city_state.get(),
                                                             entry_hiring_manager_name.get(),
                                                             is_hiring_manager_entry_enabled.get(),
                                                             is_remote_val.get(),
                                                             is_copy_to_clipboard_val.get()))

label_job_title.grid(row=0, column=0, pady=2, padx=20)
label_company_name.grid(row=1, column=0, pady=2, padx=20)
label_city_state.grid(row=2, column=0, pady=2, padx=20)

entry_job_title.grid(row=0, column=1, pady=10, padx=(0, 20))
entry_company_name.grid(row=1, column=1, pady=10, padx=(0, 20))
entry_city_state.grid(row=2, column=1, pady=10, padx=(0, 20))

check_is_hiring_manager.grid(row=4, column=1, sticky=tk.W)

label_hiring_manager_name.grid(row=5, column=0, pady=2, padx=20)
entry_hiring_manager_name.grid(row=5, column=1, pady=10, padx=(0, 20))

check_is_remote.grid(row=6, column=1, sticky=tk.W)
check_is_txt_file.grid(row=7, column=1, sticky=tk.W)

button_convert.grid(row=8, column=0, padx=(20, 20), pady=(0, 10),  columnspan=2, sticky=tk.N + tk.S + tk.W + tk.E)


def clear_values():
    entry_job_title.delete(0, tk.END)
    entry_company_name.delete(0, tk.END)
    entry_city_state.delete(0, tk.END)
    entry_hiring_manager_name.delete(0, tk.END)

    entry_job_title.focus_set()


def delete_excess_entries():
    for filename in glob.glob("C:/Users/Harrison/Documents/*"):
        if filename.__contains__(converted_file_base_name) and not filename.__contains__('_Template'):
            print(os.path.basename(filename))
            os.remove(filename)


def copy_text_from_file(key):
    data = load_data(config_file_name)

    print("Key: " + key)

    # loop config file and find by key,value pair
    for custom_button in data['customButtons']:
        if custom_button['title'] == key:
            # write the to the text file to prepare copy
            with open('output_button.txt', 'w') as fout:
                fout.writelines(custom_button['text'])

            # copy text file to clipboard
            os.system("clip < output_button.txt ")

            # delete text file
            os.remove("output_button.txt")


def add_preexisting_custom_buttons(toolbar_frame):
    if Path(config_file_name).exists():
        data = load_data(config_file_name)

        for custom_button in data['customButtons']:
            title = custom_button['title']
            print('adding button: ' + title)
            new_custom_button = tk.Button(toolbar_frame,
                                          text=title,
                                          command=lambda title=title: copy_text_from_file(title))
            new_custom_button.pack(side=tk.LEFT, padx=2, pady=5)

    # Resize the window
    window.geometry("")


# Bottom toolbar
toolbar = tk.Frame(window)
toolbar.grid(row=9, column=0, columnspan=2, padx=5, sticky=tk.N + tk.S + tk.W + tk.E)

# init toolbar
add_preexisting_custom_buttons(toolbar)

# Start menubar
menubar = tk.Menu(window)
file_menu = tk.Menu(menubar, tearoff=0)
file_menu.add_command(label="Clear Values", command=clear_values)
file_menu.add_command(label="Delete excess entries", command=delete_excess_entries)
file_menu.add_separator()
file_menu.add_command(label="Add Custom Link", command=lambda: custom_link_popup(toolbar))
menubar.add_cascade(label="File", menu=file_menu)

window.config(menu=menubar)

# calculate coordination of screen and window form
x_left = int((window.winfo_screenwidth() / 2 - tk_width / 2) * gui_start_position)
y_top = int(window.winfo_screenheight() / 2 - tk_height / 2)

# Write following format for center screen
window.geometry("+{}+{}".format(x_left, y_top))

window.mainloop()
